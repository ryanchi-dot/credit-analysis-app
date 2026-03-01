"""
Word报告生成工具
"""
import os
import json
import requests
import logging
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from langchain.tools import tool, ToolRuntime
from coze_coding_utils.runtime_ctx.context import new_context

logger = logging.getLogger(__name__)


@tool
def create_word_report(
    title: str,
    content: str,
    part_number: int = 1,
    total_parts: int = 1,
    runtime: ToolRuntime = None
) -> str:
    """
    创建授信分析报告的Word文档
    
    支持插入：
    - 标题（不同层级）
    - 文字段落
    - 表格（从Markdown表格格式转换）
    - 图片（从URL下载并插入）
    - 架构图（Mermaid或图片）
    
    Args:
        title (str): 报告标题（如："百度中国有限公司授信分析报告"）
        content (str): 报告内容（Markdown格式，包含标题、表格、图片链接等）
        part_number (int): 当前文档部分编号（用于分多个文档，默认为1）
        total_parts (int): 总共分成几个文档（默认为1）
        runtime (ToolRuntime): LangChain工具运行时上下文
    
    Returns:
        str: 生成的Word文档的URL（上传到对象存储后）
    
    示例:
        >>> create_word_report(
        ...     title="腾讯控股有限公司授信分析报告",
        ...     content="# 第一部分：申报方案分析\n## 一、多维度风险...",
        ...     part_number=1,
        ...     total_parts=2
        ... )
    """
    try:
        ctx = runtime.context if runtime else new_context(method="create_word_report")
        
        # 创建Word文档
        doc = Document()
        
        # 注释掉字体设置，避免LSP错误
        # style = doc.styles['Normal']
        # style.font.name = '宋体'
        # style.font.size = Pt(12)
        
        # 添加文档标题
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(12)
        
        # 添加部分编号（如果是多个文档）
        if total_parts > 1:
            part_para = doc.add_paragraph()
            part_run = part_para.add_run(f"【文档{part_number}/{total_parts}】")
            part_run.font.size = Pt(10)
            part_run.font.color.rgb = RGBColor(128, 128, 128)
            part_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            part_para.paragraph_format.space_after = Pt(24)
        
        # 解析内容并添加到文档
        lines = content.split('\n')
        current_table_data = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line:
                # 空行，跳过
                if current_table_data:
                    # 如果有表格数据，添加表格
                    _add_table_from_data(doc, current_table_data)
                    current_table_data = None
                continue
            
            # 检测标题
            if line.startswith('#'):
                # 如果有表格数据，先添加表格
                if current_table_data:
                    _add_table_from_data(doc, current_table_data)
                    current_table_data = None
                
                # 解析标题层级
                level = 0
                while level < len(line) and line[level] == '#':
                    level += 1
                
                heading_text = line[level:].strip()
                
                # 添加标题
                heading = doc.add_heading(heading_text, level=min(level, 3))
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 设置标题样式
                run = heading.runs[0]
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.bold = True
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)
            
            # 检测表格（Markdown格式）
            elif '|' in line and line.startswith('|'):
                if current_table_data is None:
                    current_table_data = []
                current_table_data.append(line)
            
            # 检测图片URL
            elif line.startswith('[图片') and 'http' in line:
                # 如果有表格数据，先添加表格
                if current_table_data:
                    _add_table_from_data(doc, current_table_data)
                    current_table_data = None
                
                # 提取图片URL
                url_start = line.find('http')
                url_end = line.find(')', url_start) if ')' in line[url_start:] else len(line)
                image_url = line[url_start:url_end]
                
                # 添加图片说明
                img_caption = line[line.find(':') + 1:line.find('(')].strip() if ':' in line and '(' in line else "图表"
                
                try:
                    # 下载图片
                    response = requests.get(image_url, timeout=10)
                    if response.status_code == 200:
                        # 保存到临时文件
                        temp_image_path = f"/tmp/image_{i}.png"
                        os.makedirs('/tmp', exist_ok=True)
                        with open(temp_image_path, 'wb') as f:
                            f.write(response.content)
                        
                        # 插入图片到Word
                        para = doc.add_paragraph()
                        run = para.add_run()
                        run.add_picture(temp_image_path, width=Inches(6.0))
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 添加图片说明
                        caption_para = doc.add_paragraph()
                        caption_run = caption_para.add_run(f"图表：{img_caption}")
                        caption_run.font.size = Pt(10)
                        caption_run.font.color.rgb = RGBColor(128, 128, 128)
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_para.paragraph_format.space_after = Pt(6)
                        
                        # 删除临时文件
                        if os.path.exists(temp_image_path):
                            os.remove(temp_image_path)
                except Exception as e:
                    # 图片下载失败，添加文字说明
                    para = doc.add_paragraph()
                    para.add_run(f"[图片加载失败：{img_caption}]")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 检测Mermaid图表
            elif line.startswith('```mermaid'):
                # 如果有表格数据，先添加表格
                if current_table_data:
                    _add_table_from_data(doc, current_table_data)
                    current_table_data = None
                
                # Mermaid代码块，跳过（Word不支持直接渲染）
                # 可以添加占位符
                para = doc.add_paragraph()
                para.add_run("[架构图已省略，详见原文]")
            
            # 普通段落
            else:
                # 如果有表格数据，先添加表格
                if current_table_data:
                    _add_table_from_data(doc, current_table_data)
                    current_table_data = None
                
                # 添加段落
                para = doc.add_paragraph()
                para.add_run(line)
                para.paragraph_format.space_after = Pt(6)
        
        # 如果还有表格数据，添加
        if current_table_data:
            _add_table_from_data(doc, current_table_data)
        
        # 保存Word文档到临时文件
        temp_path = f"/tmp/credit_report_part{part_number}.docx"
        os.makedirs('/tmp', exist_ok=True)
        doc.save(temp_path)
        
        # 验证文档是否正确生成
        if not os.path.exists(temp_path):
            raise ValueError("Word文档生成失败：文件不存在")
        
        file_size = os.path.getsize(temp_path)
        if file_size < 1024:  # 文件大小小于1KB，可能生成失败
            raise ValueError(f"Word文档生成失败：文件大小过小（{file_size} bytes）")
        
        logger.info(f"Word文档生成成功，文件大小：{file_size} bytes")
        
        # 读取文件内容
        with open(temp_path, 'rb') as f:
            file_content = f.read()
        
        # 使用S3Storage上传文件
        from storage.s3.s3_storage import S3SyncStorage
        import re
        
        # 获取环境变量配置
        access_key = os.getenv("COZE_WORKLOAD_IDENTITY_API_KEY", "")
        secret_key = access_key  # S3Storage 使用同一个key
        bucket_name = os.getenv("COZE_BUCKET_NAME", "")
        
        # 创建存储客户端
        storage_client = S3SyncStorage(
            access_key=access_key,
            secret_key=secret_key,
            bucket_name=bucket_name
        )
        
        # 生成文件名（只保留英文字母、数字、点、下划线、短横，其他字符替换为下划线）
        safe_title = re.sub(r'[^a-zA-Z0-9._\-]', '_', title)
        file_name = f"{safe_title}_part{part_number}.docx"
        
        logger.info(f"文件名: {file_name}, 原标题: {title}")
        
        # 上传文件
        try:
            object_key = storage_client.upload_file(
                file_content=file_content,
                file_name=file_name,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # 使用S3Storage生成签名URL（可访问的临时链接）
            try:
                # 生成1小时有效期的签名URL
                signed_url = storage_client.generate_presigned_url(
                    key=object_key,
                    expire_time=3600  # 1小时
                )
                storage_url = signed_url
                logger.info(f"生成签名URL成功: {storage_url}")
            except Exception as sign_error:
                logger.error(f"生成签名URL失败: {sign_error}")
                # 如果签名URL生成失败，回退到普通URL格式
                endpoint_url = os.getenv("COZE_BUCKET_ENDPOINT_URL", "")
                if endpoint_url:
                    clean_endpoint = endpoint_url.rstrip('/')
                    storage_url = f"{clean_endpoint}/{bucket_name}/{object_key}"
                else:
                    storage_url = f"s3://{bucket_name}/{object_key}"
            
            # 删除临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
            
            return json.dumps({
                "success": True,
                "word_url": storage_url,
                "part_number": part_number,
                "total_parts": total_parts,
                "filename": f"{title}_part{part_number}.docx",
                "object_key": object_key,
                "expires_in": "1小时" if "signed_url" in str(storage_url) else "永久"
            }, ensure_ascii=False, indent=2)
        
        except Exception as upload_error:
            logger.error(f"文件上传失败: {upload_error}")
            # 即使上传失败，也提供本地文件路径作为备用
            return json.dumps({
                "success": False,
                "error": "文件上传失败",
                "details": str(upload_error),
                "local_file": temp_path,
                "file_size": file_size
            }, ensure_ascii=False, indent=2)
    
    except Exception as e:
        import traceback
        return json.dumps({
            "success": False,
            "error": "Word文档生成失败",
            "details": str(e),
            "traceback": traceback.format_exc()
        }, ensure_ascii=False, indent=2)


def _add_table_from_data(doc, table_data):
    """
    从Markdown表格数据添加表格到Word文档
    
    Args:
        doc: Word文档对象
        table_data: Markdown表格数据列表
    """
    if not table_data or len(table_data) < 2:
        return
    
    # 解析表格数据
    rows = []
    for line in table_data:
        # 移除首尾的 |
        line = line.strip('|')
        # 分割列
        cols = [col.strip() for col in line.split('|')]
        rows.append(cols)
    
    # 检查是否有分隔行（第二行通常是分隔行）
    if len(rows) > 1 and all(c.replace('-', '').replace(':', '').replace(' ', '') == '' for c in rows[1]):
        # 移除分隔行
        rows.pop(1)
    
    if len(rows) < 2:
        return
    
    # 创建表格
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    # 填充表格数据
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            
            # 设置单元格样式
            for paragraph in cell.paragraphs:
                paragraph.space_before = Pt(3)
                paragraph.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
                    
                    # 第一行加粗
                    if i == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        # 设置背景色
                        cell.background_color = "E0E0E0"
